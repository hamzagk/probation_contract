"""
Contract Management Views
Mirrors the probation management workflow for employee contract tracking.
"""

from datetime import date, datetime, timedelta
import json
import logging

from dateutil.relativedelta import relativedelta
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.exceptions import ValidationError
from django.core.mail import EmailMultiAlternatives
from django.core.paginator import Paginator
from django.core.validators import EmailValidator
from django.conf import settings
from django.db.models import Count, Q
from django.db.models.functions import Coalesce
from django.http import JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.utils import timezone
from django.views.decorators.http import require_http_methods

from hr_portal.forms import ContractImportForm
from hr_portal.models import Department, Employee

logger = logging.getLogger(__name__)
email_validator = EmailValidator()
CONTRACT_TERM_YEAR_CHOICES = [2, 3, 4, 5]


def _contracts_queryset():
    return Employee.objects.filter(
        is_contract_record=True,
        contract_start_date__isnull=False,
    ).select_related("department").annotate(
        effective_contract_end_date=Coalesce("extended_contract_end_date", "contract_end_date")
    )


def _parse_contract_date(value):
    if not value:
        return None
    if hasattr(value, "date"):
        return value.date()
    text = str(value).strip()
    for fmt in ("%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d", "%m/%d/%Y"):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue
    return None


def _clean_text(value):
    return " ".join(str(value or "").split())


def _parse_email_list(value, field_name):
    if not value:
        return []

    emails = []
    invalid = []
    for candidate in str(value).replace(";", ",").split(","):
        email = candidate.strip()
        if not email:
            continue
        try:
            email_validator(email)
            if email not in emails:
                emails.append(email)
        except ValidationError:
            invalid.append(email)

    if invalid:
        raise ValueError(f"Invalid {field_name}: {', '.join(invalid)}")
    return emails


def _default_contract_recipient(employee):
    if employee.department and employee.department.email:
        return employee.department.email
    return getattr(settings, "HR_EMAIL", "muhammad.hamza@giki.edu.pk")


def _default_contract_email_context(employee):
    return {
        "recipient_email": _default_contract_recipient(employee),
        "pa_ps_email": "",
        "cc_emails": "",
    }


def _employee_id_for_contract(serial_number):
    serial = _clean_text(serial_number).replace(".0", "")
    return f"FC2026-{serial.zfill(2)}" if serial else ""


def _contract_status_for(expiry_date):
    if expiry_date < date.today():
        return "expired"
    if (expiry_date - date.today()).days <= 60:
        return "expiring_soon"
    return "active"


def _import_contract_rows(rows):
    created_count = 0
    updated_count = 0
    skipped_count = 0
    errors = []

    for index, row in enumerate(rows, start=1):
        serial = _clean_text(row.get("serial") or row.get("S #") or row.get("S#"))
        name = _clean_text(row.get("name") or row.get("Name"))
        designation = _clean_text(row.get("designation") or row.get("Designation"))
        department_name = _clean_text(row.get("department") or row.get("Domain/Faculty") or row.get("Faculty"))
        doj = _parse_contract_date(row.get("doj") or row.get("DOJ"))
        expiry_date = _parse_contract_date(row.get("contract_expiry") or row.get("Contract Expiry"))

        if not name or not doj or not expiry_date:
            skipped_count += 1
            continue

        try:
            department = None
            if department_name:
                department, _ = Department.objects.get_or_create(
                    name=department_name,
                    defaults={
                        "email": "muhammad.hamza@giki.edu.pk",
                        "description": f"Imported contract department for {department_name}",
                    },
                )

            employee_id = _employee_id_for_contract(serial)
            employee = None
            if employee_id:
                employee = Employee.objects.filter(
                    employee_id=employee_id,
                    is_contract_record=True,
                ).first()
            if not employee and name:
                employee = Employee.objects.filter(
                    name__iexact=name,
                    is_contract_record=True,
                ).first()

            duration = max(1, (expiry_date.year - doj.year) * 12 + expiry_date.month - doj.month)

            if employee:
                employee.name = name
                employee.designation = designation or employee.designation
                employee.department = department or employee.department
                employee.start_date = doj
                employee.contract_start_date = doj
                employee.contract_duration_months = duration
                employee.contract_end_date = expiry_date
                employee.contract_status = _contract_status_for(expiry_date)
                employee.is_contract_record = True
                employee.is_contract_extended = False
                employee.extended_contract_end_date = None
                employee.save()
                updated_count += 1
            else:
                Employee.objects.create(
                    employee_id=employee_id or f"CONTRACT-{index:03d}",
                    name=name,
                    designation=designation or "Faculty",
                    department=department,
                    start_date=doj,
                    contract_start_date=doj,
                    contract_duration_months=duration,
                    contract_end_date=expiry_date,
                    contract_status=_contract_status_for(expiry_date),
                    is_contract_record=True,
                )
                created_count += 1
        except Exception as exc:
            skipped_count += 1
            errors.append(f"Row {index}: {exc}")

    return {
        "created": created_count,
        "updated": updated_count,
        "skipped": skipped_count,
        "errors": errors,
    }


def _rows_from_docx(file_obj):
    from docx import Document

    document = Document(file_obj)
    if not document.tables:
        return []

    table = document.tables[0]
    headers = [_clean_text(cell.text) for cell in table.rows[0].cells]
    rows = []
    for table_row in table.rows[1:]:
        values = [_clean_text(cell.text) for cell in table_row.cells]
        if not any(values):
            continue
        rows.append(dict(zip(headers, values)))
    return rows


def _rows_from_excel(file_obj):
    import pandas as pd

    df = pd.read_excel(file_obj)
    rows = []
    for _, row in df.iterrows():
        rows.append({str(key): value for key, value in row.items()})
    return rows


@login_required
def contract_dashboard(request):
    today = date.today()
    employees = _contracts_queryset()
    active_base = employees.filter(contract_status__in=["active", "expiring_soon"])

    context = {
        "total_contracts": employees.count(),
        "active_count": employees.filter(contract_status="active").count(),
        "ending_soon_count": employees.filter(contract_status="expiring_soon").count(),
        "expired_count": employees.filter(contract_status="expired").count(),
        "renewed_count": employees.filter(contract_status="renewed").count(),
        "extended_count": employees.filter(is_contract_extended=True).count(),
        "ending_in_30_days": active_base.filter(effective_contract_end_date__range=[today, today + timedelta(days=30)]).count(),
        "ending_in_60_days": active_base.filter(effective_contract_end_date__range=[today, today + timedelta(days=60)]).count(),
        "overdue": active_base.filter(effective_contract_end_date__lt=today).count(),
        "dept_stats": employees.values("department__name").annotate(
            total=Count("id"),
            active=Count("id", filter=Q(contract_status="active")),
            ending_soon=Count("id", filter=Q(contract_status="expiring_soon")),
            expired=Count("id", filter=Q(contract_status="expired")),
        ).order_by("-total"),
        "upcoming_expirations": active_base.filter(
            effective_contract_end_date__range=[today, today + timedelta(days=60)]
        ).order_by("effective_contract_end_date", "name")[:15],
        "today": today,
    }
    return render(request, "hr_portal/contract_dashboard.html", context)


@login_required
def contract_list(request):
    status_filter = request.GET.get("status", "all")
    department_filter = request.GET.get("department", "")
    search_query = request.GET.get("q", "")
    date_range = request.GET.get("date_range", "")
    extension_filter = request.GET.get("extension", "")

    employees = _contracts_queryset()

    if status_filter and status_filter != "all":
        if status_filter in ["active", "expiring_soon", "expired", "renewed"]:
            employees = employees.filter(contract_status=status_filter)

    if department_filter and department_filter != "all":
        employees = employees.filter(department_id=department_filter)

    if extension_filter == "extended":
        employees = employees.filter(is_contract_extended=True)
    elif extension_filter == "not_extended":
        employees = employees.filter(is_contract_extended=False)

    today = date.today()
    if date_range == "ending_30":
        employees = employees.filter(effective_contract_end_date__range=[today, today + timedelta(days=30)])
    elif date_range == "ending_60":
        employees = employees.filter(effective_contract_end_date__range=[today, today + timedelta(days=60)])
    elif date_range == "overdue":
        employees = employees.filter(effective_contract_end_date__lt=today, contract_status__in=["active", "expiring_soon"])
    elif date_range == "expired":
        employees = employees.filter(contract_status="expired")

    if search_query:
        employees = employees.filter(
            Q(name__icontains=search_query)
            | Q(employee_id__icontains=search_query)
            | Q(designation__icontains=search_query)
            | Q(department__name__icontains=search_query)
        )

    employees = employees.order_by("effective_contract_end_date", "name")
    paginator = Paginator(employees, 20)
    page_obj = paginator.get_page(request.GET.get("page"))

    departments = Department.objects.annotate(
        employee_count=Count("employee", filter=Q(employee__is_contract_record=True))
    ).filter(
        employee_count__gt=0
    ).order_by("name")

    context = {
        "page_obj": page_obj,
        "status_filter": status_filter,
        "department_filter": department_filter,
        "search_query": search_query,
        "date_range": date_range,
        "extension_filter": extension_filter,
        "departments": departments,
        "stats": {
            "total": employees.count(),
            "active": employees.filter(contract_status="active").count(),
            "ending_soon": employees.filter(contract_status="expiring_soon").count(),
            "expired": employees.filter(contract_status="expired").count(),
            "renewed": employees.filter(contract_status="renewed").count(),
        },
        "today": today,
    }
    return render(request, "hr_portal/contract_list.html", context)


@login_required
def contract_import(request):
    if request.method == "POST":
        form = ContractImportForm(request.POST, request.FILES)
        if form.is_valid():
            upload = request.FILES["contract_file"]
            filename = upload.name.lower()
            try:
                if filename.endswith(".docx"):
                    rows = _rows_from_docx(upload)
                elif filename.endswith((".xlsx", ".xls")):
                    rows = _rows_from_excel(upload)
                else:
                    messages.error(request, "Please upload a Word (.docx) or Excel (.xlsx/.xls) file.")
                    return render(request, "hr_portal/contract_import.html", {"form": form})

                result = _import_contract_rows(rows)
                messages.success(
                    request,
                    f"Contract import complete. Created {result['created']}, updated {result['updated']}, skipped {result['skipped']}.",
                )
                for error in result["errors"][:5]:
                    messages.error(request, error)
                return redirect("contract_list")
            except Exception as exc:
                messages.error(request, f"Contract import failed: {exc}")
    else:
        form = ContractImportForm()

    return render(request, "hr_portal/contract_import.html", {"form": form})


@login_required
def contract_detail(request, employee_id):
    employee = get_object_or_404(_contracts_queryset(), pk=employee_id)
    context = {
        "employee": employee,
        "today": date.today(),
        "contract_term_year_choices": CONTRACT_TERM_YEAR_CHOICES,
        "email_form_defaults": _default_contract_email_context(employee),
    }
    return render(request, "hr_portal/contract_detail.html", context)


@login_required
@require_http_methods(["POST"])
def send_contract_renewal_email(request, employee_id):
    employee = get_object_or_404(_contracts_queryset(), pk=employee_id)
    if not (0 <= employee.days_until_contract_end <= 60):
        messages.error(request, "Contract renewal email can only be sent within 2 months before expiration.")
        return redirect("contract_detail", employee_id=employee.id)

    default_recipient = _default_contract_recipient(employee)
    try:
        recipient_email = request.POST.get("recipient_email", default_recipient).strip() or default_recipient
        email_validator(recipient_email)
        pa_ps_emails = _parse_email_list(request.POST.get("pa_ps_email", ""), "PA/PS email")
        custom_cc = _parse_email_list(request.POST.get("cc_emails", ""), "CC email")
    except (ValidationError, ValueError) as exc:
        messages.error(request, str(exc))
        return redirect("contract_detail", employee_id=employee.id)

    cc_emails = []
    for email in list(getattr(settings, "HR_EMAIL_LIST", [])) + custom_cc:
        if email and email not in cc_emails and email != recipient_email:
            cc_emails.append(email)
    to_emails = [recipient_email]
    for email in pa_ps_emails:
        if email not in to_emails:
            to_emails.append(email)

    subject = f"Contract Renewal Reminder - {employee.name} ({employee.employee_id})"
    context = {
        "employee": employee,
        "current_date": date.today(),
        "days_until_contract_end": employee.days_until_contract_end,
        "hr_contact": getattr(settings, "HR_EMAIL", "HR Department"),
        "recipient_email": recipient_email,
        "pa_ps_emails": pa_ps_emails,
    }
    html_message = render_to_string("hr_portal/contract_renewal_email.html", context)

    email = EmailMultiAlternatives(
        subject=subject,
        body=(
            f"Contract renewal reminder for {employee.name}. "
            f"Contract expires on {employee.current_contract_end_date}."
        ),
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=to_emails,
        cc=cc_emails,
    )
    email.attach_alternative(html_message, "text/html")
    email.send()

    messages.success(request, f"Contract renewal email sent to {', '.join(to_emails)}.")
    return redirect("contract_detail", employee_id=employee.id)


@login_required
@require_http_methods(["POST"])
def contract_action(request, employee_id):
    employee = get_object_or_404(_contracts_queryset(), pk=employee_id)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"success": False, "message": "Invalid JSON"})

    action = data.get("action", "").lower()
    extension_months = data.get("extension_months", 12)
    contract_start_date = _parse_contract_date(data.get("contract_start_date"))
    contract_term_years = data.get("contract_term_years")

    if action not in ["renew", "extend", "expire"]:
        return JsonResponse({"success": False, "message": "Invalid action"})

    if action == "renew":
        try:
            contract_term_years = int(contract_term_years)
        except (TypeError, ValueError):
            return JsonResponse({"success": False, "message": "Select a valid contract term between 2 and 5 years"})
        if contract_term_years not in CONTRACT_TERM_YEAR_CHOICES:
            return JsonResponse({"success": False, "message": "Select a valid contract term between 2 and 5 years"})
        if not contract_start_date:
            return JsonResponse({"success": False, "message": "Select a valid new contract start date"})

        employee.contract_start_date = contract_start_date
        employee.contract_duration_months = contract_term_years * 12
        employee.contract_end_date = employee.contract_start_date + relativedelta(months=employee.contract_duration_months)
        employee.is_contract_extended = False
        employee.extended_contract_end_date = None
        employee.contract_status = "renewed"
    elif action == "extend":
        try:
            extension_months = int(extension_months)
        except (TypeError, ValueError):
            return JsonResponse({"success": False, "message": "Invalid extension period"})
        base_end_date = employee.current_contract_end_date or employee.contract_end_date or date.today()
        employee.extended_contract_end_date = base_end_date + relativedelta(months=extension_months)
        employee.is_contract_extended = True
        employee.contract_status = "active"
    elif action == "expire":
        employee.contract_status = "expired"

    employee.save()

    return JsonResponse({
        "success": True,
        "message": f"Contract {action} action completed for {employee.name}",
        "employee_id": employee.id,
        "contract_status": employee.contract_status,
        "contract_start_date": employee.contract_start_date.strftime("%b %d, %Y") if employee.contract_start_date else "",
        "contract_end_date": employee.current_contract_end_date.isoformat() if employee.current_contract_end_date else "",
        "contract_end_date_display": employee.current_contract_end_date.strftime("%b %d, %Y") if employee.current_contract_end_date else "",
        "days_until_contract_end": employee.days_until_contract_end,
        "contract_duration_label": employee.contract_duration_label,
    })


@login_required
@require_http_methods(["POST"])
def contract_bulk_action(request):
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"success": False, "message": "Invalid JSON"})

    employee_ids = data.get("employee_ids", [])
    action = data.get("action", "")
    extension_months = data.get("extension_months", 12)

    if not employee_ids:
        return JsonResponse({"success": False, "message": "No employees selected"})

    processed_count = 0
    failed_count = 0
    for employee_id in employee_ids:
        try:
            employee = _contracts_queryset().get(pk=employee_id)
            if action == "renew":
                employee.contract_start_date = date.today()
                employee.contract_duration_months = int(extension_months)
                employee.contract_end_date = employee.contract_start_date + relativedelta(months=int(extension_months))
                employee.is_contract_extended = False
                employee.extended_contract_end_date = None
                employee.contract_status = "renewed"
            elif action == "extend":
                base_end_date = employee.current_contract_end_date or employee.contract_end_date or date.today()
                employee.extended_contract_end_date = base_end_date + relativedelta(months=int(extension_months))
                employee.is_contract_extended = True
                employee.contract_status = "active"
            elif action == "expire":
                employee.contract_status = "expired"
            else:
                failed_count += 1
                continue
            employee.save()
            processed_count += 1
        except Exception as exc:
            logger.error("Contract bulk action failed for employee %s: %s", employee_id, exc)
            failed_count += 1

    return JsonResponse({
        "success": True,
        "message": f"Processed {processed_count} contracts. {failed_count} failed.",
        "processed_count": processed_count,
        "failed_count": failed_count,
    })


@login_required
def contract_report(request):
    report_type = request.GET.get("type", "summary")
    department_id = request.GET.get("department", "")
    date_from = request.GET.get("date_from", "")
    date_to = request.GET.get("date_to", "")
    employees = _contracts_queryset()

    if department_id:
        employees = employees.filter(department_id=department_id)

    if date_from:
        try:
            employees = employees.filter(contract_start_date__gte=datetime.strptime(date_from, "%Y-%m-%d").date())
        except ValueError:
            pass

    if date_to:
        try:
            employees = employees.filter(effective_contract_end_date__lte=datetime.strptime(date_to, "%Y-%m-%d").date())
        except ValueError:
            pass

    today = date.today()
    if report_type == "detailed":
        report_data = {
            "title": "Detailed Contract Report",
            "generated_at": timezone.now(),
            "employees": employees.order_by("effective_contract_end_date", "name"),
        }
    elif report_type == "ending_soon":
        report_data = {
            "title": "Contracts Ending Soon",
            "generated_at": timezone.now(),
            "employees": employees.filter(
                effective_contract_end_date__range=[today, today + timedelta(days=60)],
                contract_status__in=["active", "expiring_soon"],
            ).order_by("effective_contract_end_date", "name"),
        }
    else:
        report_data = {
            "title": "Contract Summary Report",
            "generated_at": timezone.now(),
            "total_contracts": employees.count(),
            "by_status": {
                "active": employees.filter(contract_status="active").count(),
                "ending_soon": employees.filter(contract_status="expiring_soon").count(),
                "expired": employees.filter(contract_status="expired").count(),
                "renewed": employees.filter(contract_status="renewed").count(),
            },
            "by_timeframe": {
                "ending_30_days": employees.filter(effective_contract_end_date__range=[today, today + timedelta(days=30)]).count(),
                "ending_60_days": employees.filter(effective_contract_end_date__range=[today, today + timedelta(days=60)]).count(),
                "overdue": employees.filter(effective_contract_end_date__lt=today).count(),
            },
        }

    context = {
        "report_data": report_data,
        "report_type": report_type,
        "departments": Department.objects.order_by("name"),
        "selected_department": department_id,
        "date_from": date_from,
        "date_to": date_to,
        "today": today,
    }
    return render(request, "hr_portal/contract_report.html", context)


@login_required
def contract_statistics_api(request):
    today = date.today()
    employees = _contracts_queryset()
    return JsonResponse({
        "success": True,
        "statistics": {
            "total": employees.count(),
            "active": employees.filter(contract_status="active").count(),
            "ending_soon": employees.filter(contract_status="expiring_soon").count(),
            "expired": employees.filter(contract_status="expired").count(),
            "renewed": employees.filter(contract_status="renewed").count(),
            "ending_in_30_days": employees.filter(effective_contract_end_date__range=[today, today + timedelta(days=30)]).count(),
            "ending_in_60_days": employees.filter(effective_contract_end_date__range=[today, today + timedelta(days=60)]).count(),
            "overdue": employees.filter(effective_contract_end_date__lt=today).count(),
        },
        "generated_at": timezone.now().isoformat(),
    })


@login_required
def refresh_contract_status(request):
    updated_count = 0
    for employee in _contracts_queryset():
        old_status = employee.contract_status
        employee.save()
        if old_status != employee.contract_status:
            updated_count += 1

    return JsonResponse({
        "success": True,
        "message": f"Contract statuses refreshed. {updated_count} employees updated.",
        "updated_count": updated_count,
    })
